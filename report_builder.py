"""
report_builder.py
-----------------
Word (.docx) report generation. No Streamlit dependency.
"""
from __future__ import annotations

import io
from datetime import datetime

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor

from chart_helpers import (
    add_indicator_threshold_lines,
    build_contrib_figure,
    build_current_contributions_bar_figure,
    build_gold_signal_figure,
    build_indicator_figure,
    build_indicator_history_figure,
    build_signal_figure,
    save_fig_to_tempfile,
)
from feature_builder import classify_rsi_overlay

LEGEND_MAP = {
    "REAL_YIELD_CPI": "Real yield (10Y − CPI YoY)",
    "CPI_YOY": "Inflation (CPI YoY)",
    "USD_12M_CHG": "USD 12M % change (TWEX)",
    "CURVE_10Y_3M": "Curve (10Y–3M)",
    "DEFICIT_GDP": "Deficit % GDP",
    "REAL_YIELD_TIPS10": "Real yield (10Y TIPS, 20D MA)",
    "HY_OAS": "High Yield OAS (20D MA)",
}

_REGIME_COLORS = {
    "Structural Bull": RGBColor(0x00, 0x80, 0x00),
    "Positive": RGBColor(0x00, 0x80, 0x00),
    "Neutral": RGBColor(0x80, 0x80, 0x00),
    "Vulnerable": RGBColor(0xC0, 0x60, 0x00),
    "Structural Headwind": RGBColor(0xB0, 0x00, 0x00),
}


def _add_colored_regime_line(doc: Document, regime_label: str) -> None:
    p = doc.add_paragraph()
    p.add_run("Regime: ").bold = True
    r = p.add_run(regime_label)
    r.bold = True
    r.font.color.rgb = _REGIME_COLORS.get(regime_label, RGBColor(0x00, 0x00, 0x00))


def _simple_narrative(
    latest_row: pd.Series, core_keys: list, labels: dict, regime_label: str
) -> str:
    contribs = []
    for k in core_keys:
        c = latest_row.get(f"{k}_CONTRIB", 0.0)
        if pd.notna(c) and abs(float(c)) > 1e-9:
            contribs.append((k, float(c)))
    contribs.sort(key=lambda x: abs(x[1]), reverse=True)

    pos = next(((k, c) for k, c in contribs if c > 0), None)
    neg = next(((k, c) for k, c in contribs if c < 0), None)

    regime_simple = {
        "Structural Bull": "strongly supportive",
        "Positive": "somewhat supportive",
        "Neutral": "mixed",
        "Vulnerable": "somewhat negative",
        "Structural Headwind": "negative",
    }.get(regime_label, "mixed")

    short = {
        "REAL_YIELD_CPI": "real yields", "CPI_YOY": "inflation",
        "USD_12M_CHG": "the dollar", "CURVE_10Y_3M": "the yield curve",
        "DEFICIT_GDP": "the deficit", "REAL_YIELD_TIPS10": "TIPS real yields",
        "HY_OAS": "credit stress",
    }

    def name(k: str) -> str:
        return short.get(k, labels.get(k, k))

    if pos and neg:
        return f"Gold outlook is {regime_simple}, helped by {name(pos[0])} but hurt by {name(neg[0])}."
    if pos:
        return f"Gold outlook is {regime_simple}, mainly helped by {name(pos[0])}."
    if neg:
        return f"Gold outlook is {regime_simple}, mainly hurt by {name(neg[0])}."
    return f"Gold outlook is {regime_simple} based on the current macro mix."


def build_word_report(
    latest_row: pd.Series,
    prev_row: pd.Series,
    thresholds: dict,
    weights: dict,
    trigger_info: dict,
    labels: dict,
    legend_map: dict,
    core_keys: list,
    mode: str,
    crisis_year: str,
    res_plot: pd.DataFrame,
    include_indicator_charts: bool = True,
    gold_stats=None,
) -> bytes:
    doc = Document()
    doc.add_heading("Gold Macro Cockpit — Monthly Assessment", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Date (latest data point): {latest_row.name.date()}")

    if mode == "Structural Regime (today)":
        doc.add_paragraph("Mode: Structural Regime (today)")
    elif mode == "Crisis Similarity (template)":
        doc.add_paragraph(f"Mode: Crisis Similarity ({crisis_year} template)")
    else:
        doc.add_paragraph("Mode: Market Acceleration (fast)")

    doc.add_heading("Summary", level=2)

    if mode == "Structural Regime (today)":
        doc.add_paragraph(
            "Model Mode: Structural Regime — evaluates whether today's economic conditions "
            "are structurally supportive or hostile to gold using fixed economic thresholds."
        )
        doc.add_paragraph(
            "Regime Interpretation: "
            "Structural Bull = strong macro tailwind; Positive = mild tailwind; "
            "Neutral = balanced; Vulnerable = mild headwind; Structural Headwind = strong macro pressure."
        )
    elif mode == "Crisis Similarity (template)":
        doc.add_paragraph(
            f"Model Mode: Crisis Similarity ({crisis_year} template) — compares current "
            "macro conditions to the selected historical crisis window using quantile thresholds."
        )
        doc.add_paragraph(
            "Regime Interpretation: Structural Bull = highly similar to crisis-style "
            "gold-supportive regime; Vulnerable/Headwind = materially different from crisis regime."
        )
        if gold_stats and gold_stats.get("6m", {}).get("ok"):
            g6 = gold_stats["6m"]
            doc.add_paragraph(
                f"Crisis-conditioned gold response (within {crisis_year} window): "
                f"Among months with similar macro signal (±0.15), gold's forward {g6['horizon_m']}M return "
                f"had mean {g6['mean_ret']:.1f}%, median {g6['median_ret']:.1f}%, "
                f"and P(>0) {100 * g6['p_pos']:.0f}% (n={g6['n_samples']})."
            )
        else:
            doc.add_paragraph("Crisis-conditioned gold response: unavailable.")
    else:
        doc.add_paragraph(
            "Market Acceleration mode answers: Is the market currently rewarding gold "
            "(or pressuring it) regardless of the slow macro regime?"
        )

    if res_plot is not None and "RSI_14" in res_plot.columns:
        rsi_last = float(res_plot["RSI_14"].dropna().iloc[-1]) if res_plot["RSI_14"].dropna().shape[0] else np.nan
        rsi_z_last = float(res_plot["RSI_Z_60"].dropna().iloc[-1]) if (
            "RSI_Z_60" in res_plot.columns and res_plot["RSI_Z_60"].dropna().shape[0]) else np.nan
        rsi_slope_last = float(res_plot["RSI_SLOPE_3M"].dropna().iloc[-1]) if (
            "RSI_SLOPE_3M" in res_plot.columns and res_plot["RSI_SLOPE_3M"].dropna().shape[0]) else np.nan
        doc.add_paragraph(
            "RSI overlay (not scored): "
            f"RSI(14M)={'—' if pd.isna(rsi_last) else f'{rsi_last:.0f}'}, "
            f"RSI z-score (5Y)={'—' if pd.isna(rsi_z_last) else f'{rsi_z_last:+.2f}'}, "
            f"RSI slope (3M)={'—' if pd.isna(rsi_slope_last) else f'{rsi_slope_last:+.0f}'}. "
            f"Interpretation: {classify_rsi_overlay(rsi_last, rsi_z_last, rsi_slope_last)}"
        )

    doc.add_paragraph(
        "Signal Definition: Each indicator is scored (-1 / 0 / +1) relative to its threshold band "
        "and combined using normalized weights."
    )
    doc.add_paragraph(
        "Note: Regime classification describes macro backdrop conditions, not direct trading signals."
    )
    doc.add_paragraph("")

    sig = float(latest_row["SIGNAL"])
    sig_prev = float(prev_row["SIGNAL"])
    doc.add_paragraph(f"Weighted Signal: {sig:.2f} (Δ {sig - sig_prev:+.2f} vs previous month)")

    regime_label = str(latest_row.get("REGIME", "—"))
    _add_colored_regime_line(doc, regime_label)

    narr = _simple_narrative(latest_row, core_keys, labels, regime_label)
    doc.add_paragraph(f"Plain-language summary: {narr}")
    doc.add_paragraph(f"Weights: {', '.join([f'{k}={weights[k]:.2f}' for k in core_keys])}")

    doc.add_heading("Triggers", level=2)
    doc.add_paragraph(
        f"Bull trigger: {'ON' if trigger_info['bull_now'] else 'OFF'} "
        f"(threshold {trigger_info['trig_hi']:.2f}, persistence {trigger_info['persist']} months)"
    )
    doc.add_paragraph(
        f"Bear trigger: {'ON' if trigger_info['bear_now'] else 'OFF'} "
        f"(threshold {trigger_info['trig_lo']:.2f}, persistence {trigger_info['persist']} months)"
    )

    doc.add_heading("Model Weights", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Indicator"
    t.rows[0].cells[1].text = "Weight"
    for k, w in weights.items():
        row = t.add_row().cells
        row[0].text = legend_map.get(k, k)
        row[1].text = f"{w:.2f}"

    doc.add_heading("Indicator States (latest)", level=2)
    table = doc.add_table(rows=1, cols=5)
    h = table.rows[0].cells
    h[0].text = "Indicator"
    h[1].text = "Latest"
    h[2].text = "State (-1: bearish/0: neutral/+1: bullish)"
    h[3].text = "Contribution"
    h[4].text = "Thresholds (bull/bear)"
    for col in core_keys:
        row = table.add_row().cells
        row[0].text = labels.get(col, col)
        row[1].text = f"{float(latest_row[col]):.2f}"
        row[2].text = str(int(latest_row.get(col + "_STATE", 0)))
        row[3].text = f"{float(latest_row.get(col + '_CONTRIB', 0.0)):+.2f}"
        row[4].text = f"{thresholds[col]['bull']:.2f} / {thresholds[col]['bear']:.2f}"

    if include_indicator_charts:
        doc.add_heading("Indicator Charts (history)", level=2)
        for k in core_keys:
            fig_ind = build_indicator_history_figure(res_plot, k, labels.get(k, k), thresholds, mode)
            if fig_ind:
                imgp = save_fig_to_tempfile(fig_ind)
                doc.add_picture(imgp, width=Inches(6))

    doc.add_heading("Charts", level=2)
    fig1 = build_signal_figure(res_plot)
    img1 = save_fig_to_tempfile(fig1)
    doc.add_paragraph("Signal over time")
    doc.add_picture(img1, width=Inches(6))

    fig2 = build_gold_signal_figure(res_plot)
    if fig2:
        img2 = save_fig_to_tempfile(fig2)
        doc.add_paragraph("Gold vs Signal")
        doc.add_picture(img2, width=Inches(6))

    fig3 = build_contrib_figure(res_plot)
    if fig3:
        img3 = save_fig_to_tempfile(fig3)
        doc.add_paragraph("Contributions (stacked)")
        doc.add_picture(img3, width=Inches(6))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_word_report_compare(
    left_title: str,
    right_title: str,
    left_pack: tuple,
    right_pack: tuple,
    trigger_info: dict,
    labels_left: dict,
    labels_right: dict,
    legend_map: dict,
) -> bytes:
    (resL, coreL, thrL, dirsL, wL) = left_pack
    (resR, coreR, thrR, dirsR, wR) = right_pack

    doc = Document()
    doc.add_heading("Gold Macro Cockpit — Comparison Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    def add_mode_section(title, res_panel, core_keys, thresholds, weights, labels_dict):
        latest_row = res_panel.iloc[-1]
        prev_row = res_panel.iloc[-2] if len(res_panel) > 1 else latest_row

        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Latest data point: {latest_row.name.date()}")
        doc.add_paragraph(
            f"Signal: {float(latest_row['SIGNAL']):.2f} "
            f"(Δ {float(latest_row['SIGNAL'] - prev_row['SIGNAL']):+.2f})"
        )
        doc.add_paragraph(f"Regime: {str(latest_row.get('REGIME', '—'))}")

        doc.add_heading("Weights", level=3)
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "Indicator"
        t.rows[0].cells[1].text = "Weight"
        for k in core_keys:
            row = t.add_row().cells
            row[0].text = legend_map.get(k, labels_dict.get(k, k))
            row[1].text = f"{weights.get(k, 0.0):.2f}"

        doc.add_heading("Indicator States (latest)", level=3)
        table = doc.add_table(rows=1, cols=5)
        h = table.rows[0].cells
        for i, txt in enumerate(["Indicator", "Latest", "State", "Contribution", "Thresholds (bull/bear)"]):
            h[i].text = txt
        for col in core_keys:
            row = table.add_row().cells
            row[0].text = labels_dict.get(col, col)
            row[1].text = f"{float(latest_row[col]):.2f}" if pd.notna(latest_row.get(col)) else "—"
            row[2].text = str(int(latest_row.get(col + "_STATE", 0)))
            row[3].text = f"{float(latest_row.get(col + '_CONTRIB', 0.0)):+.2f}"
            row[4].text = f"{thresholds[col]['bull']:.2f} / {thresholds[col]['bear']:.2f}"

        doc.add_heading("Indicator Charts", level=3)
        for k in core_keys:
            figk = build_indicator_figure(res_panel, k, thresholds, title, labels_dict)
            if figk:
                imgk = save_fig_to_tempfile(figk)
                doc.add_paragraph(labels_dict.get(k, k))
                doc.add_picture(imgk, width=Inches(6))

        doc.add_heading("Charts", level=3)
        fig1 = build_signal_figure(res_panel)
        img1 = save_fig_to_tempfile(fig1)
        doc.add_paragraph("Gold Signal over time")
        doc.add_picture(img1, width=Inches(6))

        fig2 = build_gold_signal_figure(res_panel)
        if fig2:
            img2 = save_fig_to_tempfile(fig2)
            doc.add_paragraph("Gold vs Signal")
            doc.add_picture(img2, width=Inches(6))

        fig3 = build_contrib_figure(res_panel)
        if fig3:
            img3 = save_fig_to_tempfile(fig3)
            doc.add_paragraph("Contributions (stacked)")
            doc.add_picture(img3, width=Inches(6))

    add_mode_section(left_title, resL, coreL, thrL, wL, labels_left)
    doc.add_page_break()
    add_mode_section(right_title, resR, coreR, thrR, wR, labels_right)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
